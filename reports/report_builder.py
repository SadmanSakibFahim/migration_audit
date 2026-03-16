import base64
import os
import json
import hashlib
from collections import defaultdict
from datetime import date, datetime

from docx import Document
from xhtml2pdf import pisa  # type: ignore[import-untyped]
from typing import Any, Dict, List, Optional

from core.audit.enums import CheckStatus
from core.audit.verdict import final_verdict

SECTION_MAP = {
    "volume": ("Data Volume Checks", ["volume"]),
    "aggregates": (
        "Aggregate Checks",
        ["sum", "average", "avg", "max", "min", "variance"],
    ),
    "mappings": ("Mapping Checks", ["mapping"]),
    "relationships": ("Relationship Checks", ["foreign key"]),
    "data_constraints": (
        "Data Constraint Checks",
        ["data constraint", "data constraints"],
    ),
}

def get_content_hash(content: Dict[str, Any]) -> str:
    """Generate a canonical SHA-256 hash of the report content to ensure data integrity."""
    serializable = {
        "client": content["client"],
        "migration": content["migration"],
        "date": content["date"],
        "final_verdict": content["final_verdict"],
        "results": [
            {"name": r.name, "status": str(r.status.value) if hasattr(r.status, "value") else str(r.status), "message": r.message} 
            for r in content.get("all_results", [])
        ]
    }
    content_str = json.dumps(serializable, sort_keys=True)
    return hashlib.sha256(content_str.encode("utf-8")).hexdigest()

def sign_report(file_path: str) -> str:
    """Generate a .sha256 signature file for a generated report."""
    if not os.path.exists(file_path):
        return ""
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        hasher.update(f.read())
    file_hash = hasher.hexdigest()
    with open(f"{file_path}.sha256", "w") as f:
        f.write(f"{file_hash} *{os.path.basename(file_path)}\n")
    return file_hash


def group_results(results: List[Any]) -> Dict[str, Dict[str, List[Any]]]:
    """Group results hierarchically with sub-group and target extraction."""
    hierarchical = defaultdict(lambda: defaultdict(list))

    for r in results:
        name_lower = r.name.lower()
        section_assigned = "General Checks"

        for key, (section_name, keywords) in SECTION_MAP.items():
            if any(kw in name_lower for kw in keywords):
                section_assigned = section_name
                break
        
        # Determine sub-group and extract target
        sub_group = "General"
        target = r.name
        if " check: " in name_lower:
            parts = r.name.split(" Check:", 1)
            sub_group = parts[0] + " Check"
            target = parts[1].lstrip(": ").strip()
        elif ":" in r.name:
            parts = r.name.split(":", 1)
            sub_group = parts[0].strip()
            target = parts[1].strip()
        
        # Attach temporary target attribute for the writers
        r._report_target = target
        hierarchical[section_assigned][sub_group].append(r)

    return hierarchical


def section_verdict(sub_groups: Dict[str, List[Any]]) -> str:
    """Compute verdict for a section based on its sub-groups."""
    all_results = [r for sub in sub_groups.values() for r in sub]
    if any(r.status == CheckStatus.FAIL for r in all_results):
        return "FAIL"
    if any(r.status == CheckStatus.WARN for r in all_results):
        return "WARN"
    return "PASS"


def _build_report_content(results: List[Any], client: str = "Client", migration: str = "Source → Target") -> Dict[str, Any]:
    """Build the core report content as a dictionary for reuse across formats."""
    grouped = group_results(results)
    final = final_verdict(results)

    content = {
        "client": client,
        "migration": migration,
        "date": str(date.today()),
        "final_verdict": final,
        "grouped_results": grouped,
        "all_results": results,
    }
    content["integrity_hash"] = get_content_hash(content)
    return content



def _write_docx(content: Dict[str, Any], output_path: str) -> None:
    """Generate DOCX report."""
    doc = Document()

    # Title
    doc.add_heading("Migration Validation & Risk Audit", level=1)

    doc.add_paragraph(
        f"Client: {content['client']}\n"
        f"Migration: {content['migration']}\n"
        f"Audit Date: {content['date']}\n"
        f"Auditor: Albatross Audit"
    )

    # Executive Summary
    doc.add_heading("Executive Summary", level=2)

    doc.add_paragraph(f"Final Verdict: {content['final_verdict']}\n")

    for section, res in content["grouped_results"].items():
        verdict = section_verdict(res)
        doc.add_paragraph(f"{section}: {verdict}")

    for section, sub_groups in content["grouped_results"].items():
        doc.add_heading(section, level=2)

        for sub_group, res_list in sub_groups.items():
            doc.add_heading(sub_group, level=3)
            doc.add_paragraph("Findings:")
            for r in res_list:
                doc.add_paragraph(f"[{r.status.value}] {r.message}", style="List Bullet")

        doc.add_paragraph(f"Section Verdict: {section_verdict(sub_groups)}")

    # Final Verdict
    doc.add_heading("Final Deployability Verdict", level=2)
    doc.add_paragraph(content["final_verdict"])

    if "integrity_hash" in content:
        doc.add_heading("Cryptographic Signature", level=2)
        doc.add_paragraph(f"SHA-256 Data Footprint: {content['integrity_hash']}")

    doc.save(output_path)


def _write_markdown(content: Dict[str, Any], output_path: str) -> None:
    """Generate Markdown report."""
    lines = []

    lines.append("# Migration Validation & Risk Audit\n")
    lines.append(f"**Client:** {content['client']}\n")
    lines.append(f"**Migration:** {content['migration']}\n")
    lines.append(f"**Audit Date:** {content['date']}\n")
    lines.append("**Auditor:** Albatross Audit\n")

    # Executive Summary
    lines.append("\n## Executive Summary\n")
    lines.append(f"**Final Verdict:** {content['final_verdict']}\n")

    for section, res in content["grouped_results"].items():
        verdict = section_verdict(res)
        lines.append(f"- **{section}:** {verdict}")

    # Detailed Sections
    for section, sub_groups in content["grouped_results"].items():
        lines.append(f"\n## {section}\n")

        for sub_group, res_list in sub_groups.items():
            lines.append(f"### {sub_group}\n")
            for r in res_list:
                lines.append(f"- **[{r.status.value}]** {r.message}")
                if r.details:
                    lines.append(f"  - Details: {r.details}")

        lines.append(f"\n**Section Verdict:** {section_verdict(sub_groups)}\n")

    # Final Verdict
    lines.append("\n## Final Deployability Verdict\n")
    lines.append(f"{content['final_verdict']}\n")

    if "integrity_hash" in content:
        lines.append("\n## Cryptographic Signature\n")
        lines.append(f"**SHA-256 Data Footprint:** `{content['integrity_hash']}`\n")

    with open(output_path, "w") as f:
        f.write("\n".join(lines))


def _write_text(content: Dict[str, Any], output_path: str) -> None:
    """Generate plain text report."""
    lines = []

    lines.append("=" * 80)
    lines.append("MIGRATION VALIDATION & RISK AUDIT")
    lines.append("=" * 80)
    lines.append("")
    lines.append(f"Client:  {content['client']}")
    lines.append(f"Migration: {content['migration']}")
    lines.append(f"Audit Date: {content['date']}")
    lines.append("Auditor: Albatross Audit")
    lines.append("")

    # Executive Summary
    lines.append("-" * 80)
    lines.append("EXECUTIVE SUMMARY")
    lines.append("-" * 80)
    lines.append(f"Final Verdict: {content['final_verdict']}")
    lines.append("")

    for section, res in content["grouped_results"].items():
        verdict = section_verdict(res)
        lines.append(f"{section}: {verdict}")

    lines.append("")

    # Detailed Sections
    for section, sub_groups in content["grouped_results"].items():
        lines.append("-" * 80)
        lines.append(section.upper())
        lines.append("-" * 80)
        lines.append("")

        for sub_group, res_list in sub_groups.items():
            lines.append(f"  {sub_group}:")
            for r in res_list:
                lines.append(f"    [{r.status.value}] {r.message}")
                if r.details:
                    lines.append(f"         Details: {r.details}")
            lines.append("")

        lines.append(f"Section Verdict: {section_verdict(sub_groups)}")
        lines.append("")

    # Final Verdict
    lines.append("=" * 80)
    lines.append("FINAL DEPLOYABILITY VERDICT")
    lines.append("=" * 80)
    lines.append(content["final_verdict"])
    lines.append("")

    if "integrity_hash" in content:
        lines.append("-" * 80)
        lines.append("CRYPTOGRAPHIC SIGNATURE")
        lines.append("-" * 80)
        lines.append(f"SHA-256 Data Footprint: {content['integrity_hash']}")
        lines.append("")

    with open(output_path, "w") as f:
        f.write("\n".join(lines))


def _write_json(content: Dict[str, Any], output_path: str) -> None:
    """Generate JSON report."""
    serializable = {
        "metadata": {
            "client": content["client"],
            "migration": content["migration"],
            "date": content["date"],
            "auditor": "Albatross Audit",
            "integrity_hash": content.get("integrity_hash", "")
        },
        "summary": {
            "final_verdict": content["final_verdict"],
            "section_verdicts": {
                section: section_verdict(sub_groups)
                for section, sub_groups in content["grouped_results"].items()
            }
        },
        "results": [
            {
                "name": r.name,
                "status": r.status.value if hasattr(r.status, "value") else str(r.status),
                "message": r.message,
                "details": r.details,
                "metrics": r.metrics
            }
            for r in content.get("all_results", [])
        ]
    }
    with open(output_path, "w") as f:
        json.dump(serializable, f, indent=4)


def _write_html(content: Dict[str, Any], output_path: str, logo_path: Optional[str] = None, for_pdf: bool = False) -> str:
    """Generate HTML report."""

    # Handle Logo (Base64 Encode)
    logo_html = ""
    if logo_path and os.path.exists(logo_path):
        try:
            with open(logo_path, "rb") as img_file:
                b64_string = base64.b64encode(img_file.read()).decode("utf-8")
                ext = os.path.splitext(logo_path)[1].lower().replace(".", "")
                mime_type = "image/png" if ext == "png" else "image/jpeg"
                logo_html = f'<img src="data:{mime_type};base64,{b64_string}" class="logo" />'
        except Exception as e:
            print(f"Failed to load logo: {e}")

    html = f"""
    <html>
    <head>
        <style>
            @page {{
                size: A4;
                margin: 1.5cm;
                @frame footer_frame {{
                    -pdf-frame-content: footerContent;
                    bottom: 1cm;
                    margin-left: 1.5cm;
                    margin-right: 1.5cm;
                    height: 1cm;
                }}
            }}
            
            body {{ 
                font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; 
                color: #232f3e;
                line-height: 1.4;
                background-color: #fff;
            }}
            
            #footerContent {{
                font-size: 8pt;
                color: #aebbc1;
                text-align: center;
                border-top: 1px solid #f2f4f5;
                padding-top: 5px;
            }}
            
            .header-table {{
                width: 100%;
                border-bottom: 4px solid #FF7F50;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            .header-info {{ text-align: left; vertical-align: bottom; }}
            .header-logo {{ text-align: right; vertical-align: bottom; }}
            .logo {{ max-height: 50px; }}
            
            h1 {{ color: #1a202c; font-size: 26pt; margin: 0; font-weight: 800; }}
            h2 {{ 
                color: #2d3748; 
                border-left: 5px solid #FF7F50; 
                padding-left: 12px; 
                margin-top: 35px;
                font-size: 18pt;
                font-weight: 700;
                background-color: #fffaf0;
                padding-top: 8px;
                padding-bottom: 8px;
            }}
            h3 {{ color: #4a5568; font-size: 14pt; margin-top: 20px; margin-bottom: 10px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.05em; }}

            .meta-info {{ font-size: 10pt; color: #718096; margin-top: 5px; }}
            
            .verdict-box {{
                background: linear-gradient(135deg, #fffaf0 0%, #fff 100%);
                border: 1px solid #fed7d7;
                padding: 20px;
                border-radius: 12px;
                margin-bottom: 25px;
                font-size: 16pt;
                text-align: center;
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
            }}
            
            .verdict-PASS {{ color: #38a169; font-weight: 800; }}
            .verdict-WARN {{ color: #dd6b20; font-weight: 800; }}
            .verdict-FAIL {{ color: #e53e3e; font-weight: 800; }}
            .verdict-ERROR {{ color: #742a2a; font-weight: 800; }}
            
            .summary-table {{ width: 100%; border-collapse: collapse; margin-bottom: 30px; }}
            .summary-row td {{ padding: 12px; border-bottom: 1px solid #edf2f7; font-size: 11pt; }}
            .summary-row:last-child td {{ border-bottom: none; }}
            
            .detail-table {{ 
                width: 100%; 
                border-collapse: collapse; 
                margin-top: 15px; 
                font-size: 10pt;
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                overflow: hidden;
            }}
            .detail-table th {{ 
                background-color: #f7fafc; 
                color: #4a5568; 
                padding: 12px; 
                text-align: left; 
                border-bottom: 2px solid #e2e8f0;
                font-weight: 700;
            }}
            .detail-table td {{ 
                border-bottom: 1px solid #edf2f7; 
                padding: 12px; 
                vertical-align: middle;
            }}
            
            .sub-group-cell {{
                background-color: #fcfcfc;
                font-weight: 700;
                color: #2d3748;
                border-right: 1px solid #edf2f7;
                width: 25%;
            }}

            .status-badge {{
                display: inline-block;
                padding: 4px 8px;
                border-radius: 6px;
                color: white;
                font-size: 8.5pt;
                font-weight: 700;
                text-align: center;
                min-width: 60px;
                letter-spacing: 0.02em;
            }}
            .bg-PASS {{ background-color: #48bb78; text-shadow: 0 1px 1px rgba(0,0,0,0.1); }}
            .bg-WARN {{ background-color: #ed8936; text-shadow: 0 1px 1px rgba(0,0,0,0.1); }}
            .bg-FAIL {{ background-color: #f56565; text-shadow: 0 1px 1px rgba(0,0,0,0.1); }}
            .bg-ERROR {{ background-color: #9b2c2c; text-shadow: 0 1px 1px rgba(0,0,0,0.1); }}
            
            .finding-msg {{ font-weight: 500; margin-bottom: 4px; display: block; }}
            .finding-details {{ color: #718096; font-size: 9pt; font-family: 'Courier New', monospace; background: #f7fafc; padding: 2px 4px; border-radius: 3px; }}

            .section-verdict-banner {{
                margin-top: 15px;
                text-align: right;
                font-size: 10pt;
                color: #4a5568;
                font-weight: 600;
            }}
        </style>
    </head>
    <body>
        <table class="header-table">
            <tr>
                <td class="header-info">
                    <h1>Audit Intelligence</h1>
                    <div class="meta-info">
                        <strong>Target:</strong> {content['client']} · {content['migration']}<br/>
                        <strong>Date:</strong> {content['date']} · <strong>Report ID:</strong> {content.get('integrity_hash', 'N/A')[:8]}
                    </div>
                </td>
                <td class="header-logo">{logo_html}</td>
            </tr>
        </table>

        <h2>Executive Conclusion</h2>
        <div class="verdict-box">
            Deployment Readiness Score: <span class="verdict-{content['final_verdict'].replace(' ', '-')}">{content['final_verdict']}</span>
        </div>
        
        <table class="summary-table">
    """

    for section, sub_groups in content["grouped_results"].items():
        verdict = section_verdict(sub_groups)
        html += f"""
            <tr class="summary-row">
                <td><strong>{section}</strong></td>
                <td style="text-align: right;"><span class="verdict-{verdict}">{verdict}</span></td>
            </tr>
        """

    html += "</table>"

    for section, sub_groups in content["grouped_results"].items():
        html += f"""
        <div class="section-block">
            <h2>{section}</h2>
            <table class="detail-table">
                <tr>
                    <th>Check Type</th>
                    <th>Target Object</th>
                    <th>Status</th>
                    <th>Audit Findings</th>
                </tr>
        """
        for sub_group, res_list in sub_groups.items():
            rowspan = len(res_list)
            for i, r in enumerate(res_list):
                target_val = getattr(r, "_report_target", r.name)
                html += "<tr>"
                if for_pdf:
                    # xhtml2pdf doesn't support rowspan — use flat rows with visual grouping
                    top_border = "border-top: 2px solid #FF7F50;" if i == 0 else ""
                    label = sub_group if i == 0 else ""
                    html += f'<td class="sub-group-cell" style="{top_border}">{label}</td>'
                else:
                    if i == 0:
                        html += f'<td class="sub-group-cell" rowspan="{rowspan}">{sub_group}</td>'
                
                html += f"""
                    <td><code style="font-size: 9pt;">{target_val}</code></td>
                    <td style="text-align: center;">
                        <span class="status-badge bg-{r.status.value}">{r.status.value}</span>
                    </td>
                    <td>
                        <span class="finding-msg">{r.message}</span>
                        {f'<span class="finding-details">{r.details}</span>' if r.details else ''}
                    </td>
                </tr>
                """
        html += f"""
            </table>
            <div class="section-verdict-banner">Section Integrity: <span class="verdict-{section_verdict(sub_groups)}">{section_verdict(sub_groups)}</span></div>
        </div>
        """

    html += f"""
        <div style="page-break-before: always;"></div>
        <h2>Final Cryptographic Proof</h2>
        <div style="font-family: monospace; font-size: 9pt; color: #4a5568; background: #f7fafc; padding: 15px; border-radius: 8px; border: 1px solid #edf2f7; line-height: 1.6;">
            <strong>Integrity Footprint (SHA-256):</strong><br/>
            {content.get('integrity_hash', 'N/A')}
        </div>
        
        <div id="footerContent">
            Automated Audit Certification &nbsp;·&nbsp; Powered by Antigravity &nbsp;·&nbsp; Page <pdf:pagenumber> of <pdf:pagecount>
        </div>
    </body>
    </html>
    """

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
    return html


def _write_pdf(html_content: str, output_path: str) -> None:
    """Generate PDF report from HTML content using xhtml2pdf."""
    try:
        with open(output_path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(html_content, dest=pdf_file)

        if pisa_status.err:
            print(f"Error generating PDF: {pisa_status.err}")
    except Exception as e:
        print(f"Failed to generate PDF: {e}")


def build_report(
    results: List[Any],
    output_path: Optional[str] = None,
    client: str = "Client",
    migration: str = "Source → Target",
    base_dir: str = "outputs",
    label: str = "",
    logo_path: Optional[str] = None,
) -> Dict[str, str]:
    """Generate reports in all formats (DOCX, Markdown, Text, HTML, PDF).

    Args:
        results: List of TestResult objects
        output_path: Path for the DOCX file (base name). If None, generated automatically.
        client: Client name for the report
        migration: Migration description
        base_dir: Base directory for output (default: "outputs")
        label: Optional label to append to the timestamped folder (e.g., "_test")
        logo_path: Optional path to a logo image (PNG/JPG)
    """
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_client = client.replace(" ", "_").lower()
        folder_name = f"{timestamp}_{safe_client}{label}"
        output_dir = os.path.join(base_dir, folder_name)
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, "Audit_Report.docx")
    else:
        # If output_path is provided, ensure its directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Build report content once
    content = _build_report_content(results, client, migration)

    # Generate DOCX
    _write_docx(content, output_path)

    # Generate Markdown, Text, HTML, PDF, JSON
    base_path = os.path.splitext(output_path)[0]
    md_path = base_path + ".md"
    txt_path = base_path + ".txt"
    html_path = base_path + ".html"
    pdf_path = base_path + ".pdf"
    json_path = base_path + ".json"

    _write_markdown(content, md_path)
    _write_text(content, txt_path)
    # Browser HTML with full rowspan layout
    html_content = _write_html(content, html_path, logo_path=logo_path)
    # PDF-safe HTML (no rowspan — xhtml2pdf/reportlab doesn't support it)
    pdf_html = _write_html(content, html_path + ".tmp", logo_path=logo_path, for_pdf=True)
    _write_pdf(pdf_html, pdf_path)
    try:
        os.remove(html_path + ".tmp")
    except OSError:
        pass
    _write_json(content, json_path)

    paths = {
        "docx": output_path,
        "markdown": md_path,
        "text": txt_path,
        "html": html_path,
        "pdf": pdf_path,
        "json": json_path,
    }
    
    for fmt, p in paths.items():
        sign_report(p)

    return paths
